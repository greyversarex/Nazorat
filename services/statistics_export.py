from io import BytesIO
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter


def create_statistics_word_document(stats_data, title="Омори дархостҳо", date_range=None):
    """Create a Word document with statistics data."""
    doc = Document()
    
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if date_range:
        date_para = doc.add_paragraph(f"Давра: {date_range}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    generation_date = doc.add_paragraph(f"Санаи тайёр кардан: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    generation_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_heading("Омори умумӣ", level=1)
    
    summary_table = doc.add_table(rows=4, cols=2)
    summary_table.style = 'Table Grid'
    
    summary_data = [
        ("Ҳамаи дархостҳо", str(stats_data.get('total_requests', 0))),
        ("Иҷро шуд", str(stats_data.get('completed_requests', 0))),
        ("Дар тафтиш", str(stats_data.get('under_review_requests', 0))),
        ("Фоизи иҷро", f"{stats_data.get('completion_rate', 0)}%"),
    ]
    
    for i, (label, value) in enumerate(summary_data):
        summary_table.rows[i].cells[0].text = label
        summary_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    topic_stats = stats_data.get('topic_stats', [])
    if topic_stats:
        doc.add_heading("Омор аз рӯи мавзӯъҳо", level=1)
        
        topic_table = doc.add_table(rows=len(topic_stats) + 1, cols=5)
        topic_table.style = 'Table Grid'
        
        headers = ["Мавзӯъ", "Ҳамагӣ", "Иҷро шуд", "Дар тафтиш", "Фоиз"]
        for i, header in enumerate(headers):
            topic_table.rows[0].cells[i].text = header
            topic_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for row_idx, topic in enumerate(topic_stats, start=1):
            topic_table.rows[row_idx].cells[0].text = topic.get('title', '')
            topic_table.rows[row_idx].cells[1].text = str(topic.get('count', 0))
            topic_table.rows[row_idx].cells[2].text = str(topic.get('completed', 0))
            topic_table.rows[row_idx].cells[3].text = str(topic.get('pending', 0))
            topic_table.rows[row_idx].cells[4].text = f"{topic.get('percentage', 0)}%"
    
    doc.add_paragraph()
    
    if 'total_users' in stats_data:
        doc.add_heading("Омори корбарон", level=1)
        
        users_table = doc.add_table(rows=2, cols=2)
        users_table.style = 'Table Grid'
        
        users_table.rows[0].cells[0].text = "Корбарон"
        users_table.rows[0].cells[1].text = str(stats_data.get('total_users', 0))
        users_table.rows[1].cells[0].text = "Администраторҳо"
        users_table.rows[1].cells[1].text = str(stats_data.get('total_admins', 0))
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_worker_statistics_word_document(worker_data, requests_list):
    """Create a Word document with worker-specific statistics."""
    doc = Document()
    
    heading = doc.add_heading(f"Омори корбар: {worker_data.get('full_name', worker_data.get('username', 'Номаълум'))}", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    generation_date = doc.add_paragraph(f"Санаи тайёр кардан: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    generation_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    doc.add_heading("Маълумоти корбар", level=1)
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ("Номи корбар", worker_data.get('username', '')),
        ("Номи пурра", worker_data.get('full_name', '')),
        ("Нақш", "Администратор" if worker_data.get('role') == 'admin' else "Корбар"),
        ("Санаи бақайдгирӣ", worker_data.get('created_at', '')),
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = str(value)
    
    doc.add_paragraph()
    
    doc.add_heading("Омори дархостҳо", level=1)
    
    total_requests = len(requests_list)
    completed = sum(1 for r in requests_list if r.get('status') == 'completed')
    under_review = total_requests - completed
    completion_rate = round((completed / total_requests * 100), 1) if total_requests > 0 else 0
    
    stats_table = doc.add_table(rows=4, cols=2)
    stats_table.style = 'Table Grid'
    
    stats_data = [
        ("Ҳамаи дархостҳо", str(total_requests)),
        ("Иҷро шуд", str(completed)),
        ("Дар тафтиш", str(under_review)),
        ("Фоизи иҷро", f"{completion_rate}%"),
    ]
    
    for i, (label, value) in enumerate(stats_data):
        stats_table.rows[i].cells[0].text = label
        stats_table.rows[i].cells[1].text = value
    
    doc.add_paragraph()
    
    if requests_list:
        doc.add_heading("Рӯйхати дархостҳо", level=1)
        
        req_table = doc.add_table(rows=min(len(requests_list), 50) + 1, cols=5)
        req_table.style = 'Table Grid'
        
        headers = ["Рақами қайд", "Мавзӯъ", "Сана", "Ҳолат", "Шарҳ"]
        for i, header in enumerate(headers):
            req_table.rows[0].cells[i].text = header
            req_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        for row_idx, req in enumerate(requests_list[:50], start=1):
            req_table.rows[row_idx].cells[0].text = req.get('reg_number', '')
            req_table.rows[row_idx].cells[1].text = req.get('topic', '')
            req_table.rows[row_idx].cells[2].text = req.get('created_at', '')
            req_table.rows[row_idx].cells[3].text = req.get('status_label', '')
            comment = req.get('comment', '')
            req_table.rows[row_idx].cells[4].text = (comment[:50] + '...' if len(comment) > 50 else comment)
        
        if len(requests_list) > 50:
            doc.add_paragraph(f"... ва боз {len(requests_list) - 50} дархости дигар")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def create_statistics_excel_document(stats_data, title="Омори дархостҳо", date_range=None):
    """Create an Excel document with statistics data."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Омор"
    
    header_font = Font(bold=True, size=14)
    bold_font = Font(bold=True)
    center_align = Alignment(horizontal='center')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    header_fill = PatternFill(start_color="0891b2", end_color="0891b2", fill_type="solid")
    header_font_white = Font(bold=True, color="FFFFFF")
    
    ws.merge_cells('A1:E1')
    ws['A1'] = title
    ws['A1'].font = header_font
    ws['A1'].alignment = center_align
    
    if date_range:
        ws.merge_cells('A2:E2')
        ws['A2'] = f"Давра: {date_range}"
        ws['A2'].alignment = center_align
    
    ws.merge_cells('A3:E3')
    ws['A3'] = f"Санаи тайёр кардан: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    ws['A3'].alignment = center_align
    
    row = 5
    ws[f'A{row}'] = "ОМОРИ УМУМӢ"
    ws[f'A{row}'].font = bold_font
    
    row += 1
    summary_data = [
        ("Ҳамаи дархостҳо", stats_data.get('total_requests', 0)),
        ("Иҷро шуд", stats_data.get('completed_requests', 0)),
        ("Дар тафтиш", stats_data.get('under_review_requests', 0)),
        ("Фоизи иҷро", f"{stats_data.get('completion_rate', 0)}%"),
    ]
    
    for label, value in summary_data:
        ws[f'A{row}'] = label
        ws[f'B{row}'] = value
        ws[f'A{row}'].border = thin_border
        ws[f'B{row}'].border = thin_border
        row += 1
    
    row += 1
    
    topic_stats = stats_data.get('topic_stats', [])
    if topic_stats:
        ws[f'A{row}'] = "ОМОР АЗ РӮИ МАВЗӮЪҲО"
        ws[f'A{row}'].font = bold_font
        row += 1
        
        headers = ["Мавзӯъ", "Ҳамагӣ", "Иҷро шуд", "Дар тафтиш", "Фоиз"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = header_font_white
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border
        row += 1
        
        for topic in topic_stats:
            ws.cell(row=row, column=1, value=topic.get('title', '')).border = thin_border
            ws.cell(row=row, column=2, value=topic.get('count', 0)).border = thin_border
            ws.cell(row=row, column=3, value=topic.get('completed', 0)).border = thin_border
            ws.cell(row=row, column=4, value=topic.get('pending', 0)).border = thin_border
            ws.cell(row=row, column=5, value=f"{topic.get('percentage', 0)}%").border = thin_border
            row += 1
    
    for col in range(1, 6):
        ws.column_dimensions[get_column_letter(col)].width = 20
    
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def create_worker_statistics_excel_document(worker_data, requests_list):
    """Create an Excel document with worker-specific statistics."""
    wb = Workbook()
    ws = wb.active
    ws.title = "Омори корбар"
    
    header_font = Font(bold=True, size=14)
    bold_font = Font(bold=True)
    center_align = Alignment(horizontal='center')
    thin_border = Border(
        left=Side(style='thin'),
        right=Side(style='thin'),
        top=Side(style='thin'),
        bottom=Side(style='thin')
    )
    header_fill = PatternFill(start_color="0891b2", end_color="0891b2", fill_type="solid")
    header_font_white = Font(bold=True, color="FFFFFF")
    
    ws.merge_cells('A1:E1')
    ws['A1'] = f"Омори корбар: {worker_data.get('full_name', worker_data.get('username', 'Номаълум'))}"
    ws['A1'].font = header_font
    ws['A1'].alignment = center_align
    
    ws.merge_cells('A2:E2')
    ws['A2'] = f"Санаи тайёр кардан: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    ws['A2'].alignment = center_align
    
    row = 4
    ws[f'A{row}'] = "МАЪЛУМОТИ КОРБАР"
    ws[f'A{row}'].font = bold_font
    row += 1
    
    info_data = [
        ("Номи корбар", worker_data.get('username', '')),
        ("Номи пурра", worker_data.get('full_name', '')),
        ("Нақш", "Администратор" if worker_data.get('role') == 'admin' else "Корбар"),
    ]
    
    for label, value in info_data:
        ws[f'A{row}'] = label
        ws[f'B{row}'] = value
        ws[f'A{row}'].border = thin_border
        ws[f'B{row}'].border = thin_border
        row += 1
    
    row += 1
    
    total_requests = len(requests_list)
    completed = sum(1 for r in requests_list if r.get('status') == 'completed')
    under_review = total_requests - completed
    completion_rate = round((completed / total_requests * 100), 1) if total_requests > 0 else 0
    
    ws[f'A{row}'] = "ОМОРИ ДАРХОСТҲО"
    ws[f'A{row}'].font = bold_font
    row += 1
    
    stats_data = [
        ("Ҳамаи дархостҳо", total_requests),
        ("Иҷро шуд", completed),
        ("Дар тафтиш", under_review),
        ("Фоизи иҷро", f"{completion_rate}%"),
    ]
    
    for label, value in stats_data:
        ws[f'A{row}'] = label
        ws[f'B{row}'] = value
        ws[f'A{row}'].border = thin_border
        ws[f'B{row}'].border = thin_border
        row += 1
    
    row += 1
    
    if requests_list:
        ws[f'A{row}'] = "РӮЙХАТИ ДАРХОСТҲО"
        ws[f'A{row}'].font = bold_font
        row += 1
        
        headers = ["Рақами қайд", "Мавзӯъ", "Сана", "Ҳолат", "Шарҳ"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.font = header_font_white
            cell.fill = header_fill
            cell.alignment = center_align
            cell.border = thin_border
        row += 1
        
        for req in requests_list:
            ws.cell(row=row, column=1, value=req.get('reg_number', '')).border = thin_border
            ws.cell(row=row, column=2, value=req.get('topic', '')).border = thin_border
            ws.cell(row=row, column=3, value=req.get('created_at', '')).border = thin_border
            ws.cell(row=row, column=4, value=req.get('status_label', '')).border = thin_border
            comment = req.get('comment', '')
            ws.cell(row=row, column=5, value=(comment[:100] if len(comment) > 100 else comment)).border = thin_border
            row += 1
    
    for col in range(1, 6):
        ws.column_dimensions[get_column_letter(col)].width = 25
    
    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    return buffer


def create_protocol_word_document(request_data, media_path=None):
    """Create a Word document for a single protocol/request."""
    import os
    
    doc = Document()
    
    heading = doc.add_heading("ПРОТОКОЛ", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    reg_para = doc.add_paragraph()
    reg_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    reg_run = reg_para.add_run(f"№ {request_data.get('reg_number', 'Н/Д')}")
    reg_run.bold = True
    reg_run.font.size = Pt(14)
    
    if request_data.get('document_number'):
        doc_para = doc.add_paragraph()
        doc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc_run = doc_para.add_run(f"Рақами ҳуҷҷат: {request_data.get('document_number')}")
        doc_run.font.size = Pt(11)
    
    doc.add_paragraph()
    
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Table Grid'
    
    info_data = [
        ("Мавзӯъ", request_data.get('topic', '')),
        ("Корбар", request_data.get('username', '')),
        ("Санаи сохтан", request_data.get('created_at', '')),
        ("Ҳолат", request_data.get('status_label', '')),
        ("Координатҳо", request_data.get('coordinates', 'Нест')),
        ("Санаи хондан", request_data.get('admin_read_at', 'Нахонда')),
    ]
    
    for i, (label, value) in enumerate(info_data):
        cell_label = info_table.rows[i].cells[0]
        cell_value = info_table.rows[i].cells[1]
        cell_label.text = label
        cell_value.text = str(value) if value else ''
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    doc.add_paragraph()
    
    doc.add_heading("Шарҳ", level=1)
    comment = request_data.get('comment', '')
    if comment:
        doc.add_paragraph(comment)
    else:
        doc.add_paragraph("Шарҳ нест")
    
    if media_path and os.path.exists(media_path):
        ext = os.path.splitext(media_path)[1].lower()
        image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']
        if ext in image_extensions:
            doc.add_paragraph()
            doc.add_heading("Расм", level=1)
            try:
                from PIL import Image
                img = Image.open(media_path)
                if img.mode in ('RGBA', 'LA', 'P'):
                    img = img.convert('RGB')
                
                img_buffer = BytesIO()
                img.save(img_buffer, format='JPEG', quality=90)
                img_buffer.seek(0)
                
                doc.add_picture(img_buffer, width=Inches(5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                doc.add_paragraph(f"Расмро илова кардан имконнопазир: {str(e)}")
        else:
            doc.add_paragraph()
            doc.add_heading("Файли замима", level=1)
            doc.add_paragraph(f"Номи файл: {os.path.basename(media_path)}")
    
    if request_data.get('admin_reply'):
        doc.add_paragraph()
        doc.add_heading("Ҷавоби админ", level=1)
        doc.add_paragraph(request_data.get('admin_reply'))
        if request_data.get('admin_reply_at'):
            reply_date = doc.add_paragraph(f"Санаи ҷавоб: {request_data.get('admin_reply_at')}")
            reply_date.runs[0].italic = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    footer = doc.add_paragraph(f"Санаи тайёр кардан: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
